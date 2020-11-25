# gui import
import tkinter as tk
from tkinter import filedialog
from tkinter import Text

# extracting text import
import os
import re
from shutil import copy
import docx
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
import csv
# Extract text from docx file


def getDocxText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    original_text = '\n'.join(fullText)
    filter_text = "\n".join([ll.rstrip() for ll in original_text.splitlines() if
                             ll.strip()])
    return filter_text.lower()

# Extract text from PDF files.


def getPDFText(filename):
    output_string = StringIO()
    with open(filename, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

    original_text = output_string.getvalue()
    filter_text = "\n".join([ll.rstrip() for ll in original_text.splitlines() if
                             ll.strip()])

    return filter_text.lower()

# get keyword count


def getCount(filter_text, keyword):
    return sum(filter_text.count(x) for x in keyword if x != '')


# Notes
#
# To open a pdf with adobe you need pass an argument of pdfpath to .exe file
# like below:
# <Acrobat path> /A "<parameter>=<value>" "<PDF path>"
# Reference: https://www.adobe.com/content/dam/acom/en/devnet/acrobat/pdfs/pdf_open_parameters.pdf
# for this we need to use os.subprocess(exepath, arguments)
# canvas size
HEIGHT = 600
WIDTH = 600

root = tk.Tk()


# canvas size
#
canvas = tk.Canvas(root, height=HEIGHT, width=WIDTH)
canvas.pack()

frame = tk.Frame(root, bg="#80c1ff", padx=20, pady=20)
frame.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.9)
e = tk.Entry(frame, width=50)
e.pack()

dirname = ""


def callback():
    global dirname
    dirname = filedialog.askdirectory()

filtered_obj = {}
def recursive_read():
    ranks = []
    global dirname
    for root, _, files in os.walk(dirname):
        for docs in files:
            text = get_text(os.path.join(root, docs))
            keyword = re.sub(' +', ',', e.get().lower())
            tup = tuple(map(str, keyword.split(',')))
            count = getCount(text, tup)
            ranks.append({'name': os.path.join(root, docs), 'count': count})

    ranks.sort(key=lambda i: i['count'], reverse=True)
    global filtered_obj
    filtered_obj = ranks
    src = ranks[0]['name']
    basename = os.path.basename(src)
    if not os.path.exists(os.path.join(dirname,'selected')):
        os.makedirs(os.path.join(dirname,'selected'))
    
    copy(src,os.path.join(dirname, 'selected', basename))
    tk.Label(frame, text='{}'.format(ranks[0]['name'])).pack()


def get_text(filename):
    if filename.lower().endswith('.pdf'):
        text = getPDFText(filename)
        return text
    elif filename.lower().endswith('.docx'):
        text = getDocxText(filename)
        return text
    else:
        return ""


def download_csv():
    global filtered_obj
    with open('filtered_resume.csv', mode='w') as csv_file:
        fieldnames = ['name','count']
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames)

        writer.writeheader()
        for obj in filtered_obj:
            writer.writerow(obj)


button = tk.Button(frame, text="Open directory", command=callback)
button.pack()

buttonCommit = tk.Button(frame, height=1, width=10, text="Print count",
                         command=recursive_read)
# command=lambda: retrieve_input() >>> just means do this when i press the button
buttonCommit.pack()

csvbutton = tk.Button(frame, text="Downalod csv", command=download_csv)
csvbutton.pack()

root.mainloop()
