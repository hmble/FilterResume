import docx
from io import StringIO

from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser


def getDocxText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def getPDFText(filename):
    output_string = StringIO()
    with open(filename, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

    return output_string.getvalue()


original_text = getPDFText('./NikhilResume.pdf')
filter_text = "\n".join([ll.rstrip() for ll in original_text.splitlines() if
                         ll.strip()])


def getCount(original_text, keyword):
    filter_text = "\n".join([ll.rstrip() for ll in original_text.splitlines() if
                             ll.strip()])
    return filter_text.count(keyword)


print(getCount(getPDFText('./NikhilResume.pdf'), "Java"))
